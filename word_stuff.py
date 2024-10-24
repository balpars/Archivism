from docx import Document
from pathlib import Path
import os

def save_transcription_as_word(transcription_result, output_folder: str, file_name: str):
    """Save the transcription result as a Word (.docx) file."""
    try:
        # Create a new Word document
        doc = Document()
        doc.add_heading('Transcription', 0)

        # Add the segments to the document
        for i, segment in enumerate(transcription_result['segments']):
            start = format_timestamp(segment['start'])
            end = format_timestamp(segment['end'])
            doc.add_paragraph(f"{i + 1}\n{start} --> {end}\n{segment['text']}\n")

        # Create output file path
        word_file = Path(output_folder) / f"{file_name}.docx"

        # Save the document
        doc.save(word_file)
        print(f"Transcription saved as Word document: {word_file}")

    except Exception as e:
        print(f"Error saving transcription as Word: {e}")
        raise

def format_timestamp(seconds: float) -> str:
    """Convert seconds to a timestamp (HH:MM:SS,mmm)."""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    milliseconds = int((seconds - int(seconds)) * 1000)
    return f"{hours:02}:{minutes:02}:{secs:02},{milliseconds:03}"

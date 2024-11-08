import os
from pathlib import Path
import time
import torch
from PyQt5.QtCore import QThread, pyqtSignal
from downloader import download_media
import subprocess
import re
from docx import Document

def create_transcription_directory(video_path):
    """
    Creates a directory for transcription output based on the video file name.
    Args:
        video_path (str): The path to the video file.
    Returns:
        str: The path to the created transcription directory.
    """
    # Get the parent directory of the video file
    video_parent_dir = os.path.dirname(video_path)

    # Get the video title and sanitize it
    video_title = os.path.splitext(os.path.basename(video_path))[0]
    sanitized_title = sanitize_title(video_title)

    # Create the transcription directory path in the same location as the video
    transcription_dir = os.path.join(video_parent_dir, f"transcription_{sanitized_title}")

    # Create the directory if it doesn't exist
    os.makedirs(transcription_dir, exist_ok=True)
    return transcription_dir

def sanitize_title(title):
    """
    Sanitizes the title to create a filesystem-friendly folder name.
    """
    # Replace any character that is not alphanumeric, space, or hyphen with an underscore
    sanitized = re.sub(r'[^\w\s-]', '', title)
    # Replace spaces and consecutive underscores with a single underscore
    sanitized = re.sub(r'\s+', '_', sanitized)
    return sanitized.strip('_')


def replace_newlines(text):
    """
    # Replace newline characters not preceded by a period, question mark, or exclamation point with a space
    """
    return re.sub(r'(?<![.!?])\n', ' ', text)

def remove_newlines(text):
    return re.sub(r'[\r\n]+', ' ', text)

class WorkerThread(QThread):
    update_label = pyqtSignal(str)
    process_finished = pyqtSignal()

    def word_routine(self, video_title, duration, output_dir):
        """
        Creates a Word document with the video title, transcription duration,
        and appends the content of the transcription text file.

        Args:
            video_title (str): The title of the video (without extension).
            duration (float): The duration of the transcription in seconds.
            output_dir (str): The directory where the Word document will be saved.
        """
        # Create a new Word document
        doc = Document()

        # Set the title of the document
        doc.add_heading(video_title, level=1)

        # Write the duration at the beginning of the document
        doc.add_paragraph(f"Transcription Duration: {duration:.2f} seconds")
        doc.add_paragraph(f"Device used: {self.device}")
        # Read the transcription text file (assuming there is only one)
        txt_files = [f for f in os.listdir(output_dir) if f.endswith('.txt')]

        if txt_files:  # Check if there is a .txt file
            txt_file_path = os.path.join(output_dir, txt_files[0])  # Get the first txt file
            with open(txt_file_path, 'r', encoding='utf-8') as txt_file:
                transcription_content = txt_file.read()
                # Append the transcription content to the Word document

                if self.noNewLine:
                    transcription_content = remove_newlines(transcription_content)
                else:
                    transcription_content = replace_newlines(transcription_content)

                doc.add_paragraph(transcription_content)

        # Save the document with the video title as the filename
        word_filename = os.path.join(output_dir, f"{video_title}.docx")
        doc.save(word_filename)

    def download_routine(self):
        for url in self.url_list:
            try:
                self.update_label.emit(f"Downloading {url}...")
                self.downloaded_list += download_media(url, self.output_folder)
                self.url_list.remove(url)  # Remove URL after processing
            except Exception as e:
                print(f"Failed to download media from {url}: {e}")

    def transcribe_routine(self):

        if torch.cuda.is_available():
            self.device = "cuda"
        else:
            self.device = "cpu"

        merge_list = self.local_list + self.downloaded_list
        path_list = []
        for f in merge_list:
            path_list.append(Path(f))

        for p in path_list:
            # Get the video title
            video_title = os.path.splitext(os.path.basename(p))[0]
            output_dir = create_transcription_directory(p)
            self.update_label.emit(f"Transcribing {video_title}...")
            # Record the start time
            start_time = time.time()
            subprocess.run(["whisper", p, "--model", "turbo", "--device", self.device,"--output_dir", output_dir])
            # Record the end time
            end_time = time.time()
            # Calculate the duration
            duration = end_time - start_time
            # Create a Word file using the word_routine function
            self.word_routine(video_title, duration, output_dir)

    def __init__(self, url_list, local_list, output_folder, mode, noNewLine):
        """ Mode = 1 Sadece indir
            Mode = 0 Transkript

            noNewLine = True or False
        """
        super().__init__()
        self.url_list = url_list
        self.local_list = local_list
        self.downloaded_list= []
        self.output_folder = output_folder
        self.mode = mode
        self.noNewLine = noNewLine
        self.device = None

    def run(self):
        """Run the download and transcribe processes."""
        if len(self.url_list) == 0 and len(self.local_list) == 0:
            self.process_finished.emit()
            return

        if self.mode == 1:  # Download only
            if len(self.url_list) > 0:
                self.download_routine()
        else:
            if len(self.url_list) > 0:
                downloaded_list = self.download_routine()

            self.transcribe_routine()
        self.process_finished.emit()
